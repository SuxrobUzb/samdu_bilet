import sys
if sys.platform == "win32":
    import pythoncom  # Faqat Windows uchun
from docx.shared import Pt  
import os
import requests
import json
from rest_framework.views import APIView
from rest_framework.response import Response
from docx import Document
from docx2pdf import convert
from PyPDF2 import PdfMerger
import random
from django.conf import settings
from rest_framework.parsers import MultiPartParser
from openpyxl import load_workbook
from rest_framework.permissions import AllowAny
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from rest_framework.parsers import MultiPartParser, FormParser
from drf_yasg.utils import swagger_auto_schema
from rest_framework_simplejwt.authentication import JWTAuthentication
from .models import SamDUkf, Uquv_yili, Bosqich, Talim_yunalishi, Semestr, Fan, SamDUkfDoc
from .serializers import (
    SamDUkfDocSerializer, SamDUkfSerializer,  UquvYiliSerializer, BosqichSerializer,
    TalimYunalishiSerializer, SemestrSerializer, FanSerializer
)
from rest_framework.viewsets import ModelViewSet
from rest_framework.permissions import IsAuthenticated
from drf_yasg import openapi
from django.core.files.base import ContentFile
from django.core.files.storage import default_storage
from rest_framework import status
from django.core.files import File



class SamDUkfDocViewSet(ModelViewSet):
    queryset = SamDUkfDoc.objects.all()
    serializer_class = SamDUkfDocSerializer

    def perform_create(self, serializer):
        # Yangi hujjat yaratilganda qo'shimcha logika qo'shish mumkin
        serializer.save() 

class UquvYiliViewSet(ModelViewSet):
    queryset = Uquv_yili.objects.all()
    serializer_class = UquvYiliSerializer
    permission_classes = [IsAuthenticated]

    @swagger_auto_schema(
        operation_description="Yangi o'quv yili qo'shish",
        request_body=UquvYiliSerializer,
        responses={
            201: UquvYiliSerializer,
            400: "Xato so'rov",
            401: "Autentifikatsiya talab qilinadi"
        }
    )
    def create(self, request, *args, **kwargs):
        return super().create(request, *args, **kwargs)

    @swagger_auto_schema(
        operation_description="O'quv yillari ro'yxatini olish",
        responses={
            200: UquvYiliSerializer(many=True),
            401: "Autentifikatsiya talab qilinadi"
        }
    )
    def list(self, request, *args, **kwargs):
        return super().list(request, *args, **kwargs)

class BosqichViewSet(ModelViewSet):
    queryset = Bosqich.objects.all()
    serializer_class = BosqichSerializer
    permission_classes = [IsAuthenticated]

    @swagger_auto_schema(
        operation_description="Yangi bosqich qo'shish",
        request_body=BosqichSerializer,
        responses={
            201: BosqichSerializer,
            400: "Xato so'rov",
            401: "Autentifikatsiya talab qilinadi"
        }
    )
    def create(self, request, *args, **kwargs):
        return super().create(request, *args, **kwargs)

    @swagger_auto_schema(
        operation_description="Bosqichlar ro'yxatini olish",
        responses={
            200: BosqichSerializer(many=True),
            401: "Autentifikatsiya talab qilinadi"
        }
    )
    def list(self, request, *args, **kwargs):
        return super().list(request, *args, **kwargs)

class TalimYunalishiViewSet(ModelViewSet):
    queryset = Talim_yunalishi.objects.all()
    serializer_class = TalimYunalishiSerializer
    permission_classes = [IsAuthenticated]

    @swagger_auto_schema(
        operation_description="Yangi ta'lim yo'nalishi qo'shish",
        request_body=TalimYunalishiSerializer,
        responses={
            201: TalimYunalishiSerializer,
            400: "Xato so'rov",
            401: "Autentifikatsiya talab qilinadi"
        }
    )
    def create(self, request, *args, **kwargs):
        return super().create(request, *args, **kwargs)

    @swagger_auto_schema(
        operation_description="Ta'lim yo'nalishlari ro'yxatini olish",
        responses={
            200: TalimYunalishiSerializer(many=True),
            401: "Autentifikatsiya talab qilinadi"
        }
    )
    def list(self, request, *args, **kwargs):
        return super().list(request, *args, **kwargs)

class SemestrViewSet(ModelViewSet):
    queryset = Semestr.objects.all()
    serializer_class = SemestrSerializer
    permission_classes = [IsAuthenticated]

    @swagger_auto_schema(
        operation_description="Yangi semestr qo'shish",
        request_body=SemestrSerializer,
        responses={
            201: SemestrSerializer,
            400: "Xato so'rov",
            401: "Autentifikatsiya talab qilinadi"
        }
    )
    def create(self, request, *args, **kwargs):
        return super().create(request, *args, **kwargs)

    @swagger_auto_schema(
        operation_description="Semestrlar ro'yxatini olish",
        responses={
            200: SemestrSerializer(many=True),
            401: "Autentifikatsiya talab qilinadi"
        }
    )
    def list(self, request, *args, **kwargs):
        return super().list(request, *args, **kwargs)

class FanViewSet(ModelViewSet):
    queryset = Fan.objects.all()
    serializer_class = FanSerializer
    permission_classes = [IsAuthenticated]

    @swagger_auto_schema(
        operation_description="Yangi fan qo'shish",
        request_body=FanSerializer,
        responses={
            201: FanSerializer,
            400: "Xato so'rov",
            401: "Autentifikatsiya talab qilinadi"
        }
    )
    def create(self, request, *args, **kwargs):
        return super().create(request, *args, **kwargs)

    @swagger_auto_schema(
        operation_description="Fanlar ro'yxatini olish",
        responses={
            200: FanSerializer(many=True),
            401: "Autentifikatsiya talab qilinadi"
        }
    )
    def list(self, request, *args, **kwargs):
        return super().list(request, *args, **kwargs)

class SamDUkfViewSet(ModelViewSet):
    queryset = SamDUkf.objects.all()
    serializer_class = SamDUkfSerializer
    permission_classes = [AllowAny]
    parser_classes = [MultiPartParser, FormParser]

    @swagger_auto_schema(
        operation_description="Yangi SamDUkf yozuvi qo'shish",
        manual_parameters=[
            openapi.Parameter(
                'uquv_yili_id', openapi.IN_FORM, type=openapi.TYPE_INTEGER,
                description='O‘quv yili ID', required=True
            ),
            openapi.Parameter(
                'semestr_id', openapi.IN_FORM, type=openapi.TYPE_INTEGER,
                description='Semestr ID', required=True
            ),
            openapi.Parameter(
                'fan_id', openapi.IN_FORM, type=openapi.TYPE_INTEGER,
                description='Fan ID', required=True
            ),
            openapi.Parameter(
                'bosqich_id', openapi.IN_FORM, type=openapi.TYPE_INTEGER,
                description='Bosqich ID', required=True
            ),
            openapi.Parameter(
                'talim_yunalishi_id', openapi.IN_FORM, type=openapi.TYPE_INTEGER,
                description='Ta‘lim yo‘nalishi ID', required=True
            ),
            openapi.Parameter(
                'file', openapi.IN_FORM, type=openapi.TYPE_FILE,
                description='Fayl (PDF, Excel yoki boshqa format)', required=False
            ),
            openapi.Parameter(
                'biletlar_soni', openapi.IN_FORM, type=openapi.TYPE_INTEGER,
                description='Biletlar soni', required=False
            ),
            openapi.Parameter(
                'oson_savol', openapi.IN_FORM, type=openapi.TYPE_INTEGER,
                description='Oson savollar soni (1-5)', required=False
            ),
            openapi.Parameter(
                'urtacha_savol', openapi.IN_FORM, type=openapi.TYPE_INTEGER,
                description='O‘rtacha savollar soni (1-5)', required=False
            ),
            openapi.Parameter(
                'murakkab1', openapi.IN_FORM, type=openapi.TYPE_INTEGER,
                description='Murakkab1 savollar soni (1-5)', required=False
            ),
            openapi.Parameter(
                'murakkab2', openapi.IN_FORM, type=openapi.TYPE_INTEGER,
                description='Murakkab2 savollar soni (1-5)', required=False
            ),
            openapi.Parameter(
                'qiyin_savol', openapi.IN_FORM, type=openapi.TYPE_INTEGER,
                description='Qiyin savollar soni (1-5)', required=False
            ),
        ],
        responses={
            201: SamDUkfSerializer,
            400: "Xato so'rov",
            401: "Autentifikatsiya talab qilinadi"
        }
    )
    def create(self, request, *args, **kwargs):
        """
        Yangi SamDUkf yozuvini yaratish.
        Agar berilgan ma'lumotlar bazada allaqachon mavjud bo'lsa,
        eski yozuvni ID orqali aniqlab olib o'chirib tashlaydi.
        Shuningdek, SamDUkfDoc modelidagi bog'liq yozuvlarni ham o'chiradi.
        """
        # Kirish ma'lumotlarini validatsiya qilish
        serializer = self.get_serializer(data=request.data)
        serializer.is_valid(raise_exception=True)

        # Bazada allaqachon mavjud bo'lgan yozuvni topish
        uquv_yili_id = request.data.get('uquv_yili_id')
        semestr_id = request.data.get('semestr_id')
        fan_id = request.data.get('fan_id')
        bosqich_id = request.data.get('bosqich_id')
        talim_yunalishi_id = request.data.get('talim_yunalishi_id')

        existing_record = SamDUkf.objects.filter(
            uquv_yili_id=uquv_yili_id,
            semestr_id=semestr_id,
            fan_id=fan_id,
            bosqich_id=bosqich_id,
            talim_yunalishi_id=talim_yunalishi_id
        ).first()

        if existing_record:
            # Eski yozuvni o'chirish
            if hasattr(existing_record, 'samdukfdoc'):
                # SamDUkfDoc modelidagi bog'liq yozuvni o'chirish
                existing_record.samdukfdoc.delete()
            existing_record.delete()

        # Yangi yozuvni saqlash
        self.perform_create(serializer)
        headers = self.get_success_headers(serializer.data)
        return Response(serializer.data, status=status.HTTP_201_CREATED, headers=headers)

    def perform_create(self, serializer):
        """
        Faylni saqlash va bog'liq modellarni yaratish.
        """
        serializer.save()


class UploadQuestions(APIView):
    parser_classes = [MultiPartParser]
    permission_classes = [AllowAny]

    def post(self, request):
        try:
            latest_entry = SamDUkf.objects.latest('id')
        except SamDUkf.DoesNotExist:
            return Response({"error": "Bazada hech qanday ma'lumot yo'q."}, status=400)

        file = latest_entry.file.file
        full_path = os.path.join(settings.MEDIA_ROOT, file.name)

        if not os.path.exists(full_path):
            return Response({"error": f"Fayl topilmadi: {full_path}"}, status=400)

        num_tickets = latest_entry.biletlar_soni
        num_easy = latest_entry.oson_savol
        num_medium = latest_entry.urtacha_savol
        num_murakkab1 = latest_entry.murakkab1
        num_murakkab2 = latest_entry.murakkab2
        num_hard = latest_entry.qiyin_savol

        for num, name in [(num_easy, "oson"), (num_medium, "o'rtacha"), (num_murakkab1, "murakkab1"), 
                          (num_murakkab2, "murakkab2"), (num_hard, "qiyin")]:
            if not (1 <= num <= 5):
                return Response({"error": f"{name} savol soni 1 dan 5 gacha bo'lishi kerak!"}, status=400)

        try:
            # Fayl turini aniqlash
            file_extension = os.path.splitext(full_path)[1].lower()
            questions_easy = []
            questions_medium = []
            questions_murakkab1 = []
            questions_murakkab2 = []
            questions_hard = []
            column_map = {1: 1, 2: 2, 3: 3, 4: 4, 5: 5}

            if file_extension == '.xlsx' or file_extension == '.xls':
                # Excel faylini o'qish
                workbook = load_workbook(full_path)
                sheet = workbook.active

                for row in sheet.iter_rows(min_row=2, min_col=1, values_only=True):
                    if len(row) >= 5:
                        if num_easy and row[column_map[num_easy]]:
                            questions_easy.append(row[column_map[num_easy]])
                        if num_medium and row[column_map[num_medium]]:
                            questions_medium.append(row[column_map[num_medium]])
                        if num_murakkab1 and row[column_map[num_murakkab1]]:
                            questions_murakkab1.append(row[column_map[num_murakkab1]])
                        if num_murakkab2 and row[column_map[num_murakkab2]]:
                            questions_murakkab2.append(row[column_map[num_murakkab2]])
                        if num_hard and row[column_map[num_hard]]:
                            questions_hard.append(row[column_map[num_hard]])

            elif file_extension == '.docx':
                # Word faylini o'qish
                doc = Document(full_path)
                table = doc.tables[0]  # Birinchi jadvalni olamiz
                for row in table.rows[1:]:  # Birinchi ikki qatorni tashlaymiz
                    cells = [cell.text.strip() for cell in row.cells][2:]  # Birinchi ikki ustunni tashlaymiz
                    if len(cells) >= 5:
                        if num_easy and cells[num_easy - 1]:  # 0 dan boshlash uchun -1
                            questions_easy.append(cells[num_easy - 1])
                        if num_medium and cells[num_medium - 1]:
                            questions_medium.append(cells[num_medium - 1])
                        if num_murakkab1 and cells[num_murakkab1 - 1]:
                            questions_murakkab1.append(cells[num_murakkab1 - 1])
                        if num_murakkab2 and cells[num_murakkab2 - 1]:
                            questions_murakkab2.append(cells[num_murakkab2 - 1])
                        if num_hard and cells[num_hard - 1]:
                            questions_hard.append(cells[num_hard - 1])
            else:
                return Response({"error": "Qo'llab-quvvatlanmaydigan fayl formati. Faqat .xlsx, .xls yoki .docx ruxsat etiladi."}, status=400)
            json_files = {
                "easy": questions_easy,
                "medium": questions_medium,
                "murakkab1": questions_murakkab1,
                "murakkab2": questions_murakkab2,
                "hard": questions_hard,
            }

            for difficulty, questions in json_files.items():
                output_path = os.path.join(settings.MEDIA_ROOT, f"{difficulty}_questions.json")
                with open(output_path, "w", encoding="utf-8") as json_file:
                    json.dump(questions, json_file, ensure_ascii=False, indent=4)

        except Exception as e:
            return Response({"error": f"Faylni o'qishda xatolik: {str(e)}"}, status=500)

        # 🎟 Biletlarni yaratish uchun API'ga so‘rov yuborish
        base_url = "http://127.0.0.1:8000"
        generate_url = f"{base_url}/api/generate_tickets/"
        export_url = f"{base_url}/api/export_tickets/"

        response = requests.post(
            generate_url,
            json={
                "num_tickets": num_tickets,
                "num_easy": num_easy,
                "num_medium": num_medium,
                "num_murakkab1": num_murakkab1,
                "num_murakkab2": num_murakkab2,
                "num_hard": num_hard
            },
        )

        if response.status_code != 200:
            return Response({"error": f"GenerateTickets xatosi: {response.text}"}, status=response.status_code)

        export_response = requests.get(export_url)

        if export_response.status_code != 200:
            return Response({"error": f"ExportTickets xatosi: {export_response.text}"}, status=export_response.status_code)

        return Response({
            "message": f"Savollar yuklandi va {num_tickets} ta bilet yaratildi!",
            "files": [
                f"{settings.MEDIA_URL}easy_questions.json",
                f"{settings.MEDIA_URL}medium_questions.json",
                f"{settings.MEDIA_URL}murakkab1_questions.json",
                f"{settings.MEDIA_URL}murakkab2_questions.json",
                f"{settings.MEDIA_URL}hard_questions.json",
            ]
        })


class GenerateTickets(APIView):
    permission_classes = [AllowAny]
    

    def post(self, request):
        # Biletlar sonini olish
        num_tickets = request.data.get('num_tickets', 1)
        try:
            num_tickets = int(num_tickets)
            if num_tickets < 1:
                return Response({"error": "Biletlar soni kamida 1 ta bo'lishi kerak!"}, status=400)
        except (ValueError, TypeError):
            return Response({"error": "num_tickets must be an integer"}, status=400)

        # Har bir darajadan savollar sonini olish (faqat JSON fayllardan 1 ta olinadi)
        num_easy = request.data.get('num_easy', 1)
        num_medium = request.data.get('num_medium', 1)
        num_murakkab1 = request.data.get('num_murakkab1', 1)
        num_murakkab2 = request.data.get('num_murakkab2', 1)
        num_hard = request.data.get('num_hard', 1)

        # Har bir daraja uchun 1-5 oralig'ini tekshirish
        for num, name in [(num_easy, "oson"), (num_medium, "o'rtacha"), (num_murakkab1, "murakkab1"), 
                          (num_murakkab2, "murakkab2"), (num_hard, "qiyin")]:
            if not (1 <= num <= 5):
                return Response({"error": f"{name} savol soni 1 dan 5 gacha bo'lishi kerak!"}, status=400)

        # JSON fayllarni o'qish
        try:
            question_files = {
                "easy": os.path.join(settings.MEDIA_ROOT, "easy_questions.json"),
                "medium": os.path.join(settings.MEDIA_ROOT, "medium_questions.json"),
                "murakkab1": os.path.join(settings.MEDIA_ROOT, "murakkab1_questions.json"),
                "murakkab2": os.path.join(settings.MEDIA_ROOT, "murakkab2_questions.json"),
                "hard": os.path.join(settings.MEDIA_ROOT, "hard_questions.json"),
            }

            questions = {}
            for level, path in question_files.items():
                with open(path, "r", encoding="utf-8") as f:
                    questions[level] = json.load(f)

            # Har bir darajada savollar borligini tekshirish
            if any(not q for q in questions.values()):
                return Response({"error": "Tanlash uchun yetarli savollar mavjud emas!"}, status=400)

        except FileNotFoundError:
            return Response({"error": "JSON fayllari topilmadi! Avval savollarni yuklang."}, status=400)

        # Biletlarni yaratish
        tickets = []
        try:
            for _ in range(num_tickets):
                ticket = {
                    "easy": random.choice(questions["easy"]) if questions["easy"] else "Savol yo'q",
                    "medium": random.choice(questions["medium"]) if questions["medium"] else "Savol yo'q",
                    "murakkab1": random.choice(questions["murakkab1"]) if questions["murakkab1"] else "Savol yo'q",
                    "murakkab2": random.choice(questions["murakkab2"]) if questions["murakkab2"] else "Savol yo'q",
                    "hard": random.choice(questions["hard"]) if questions["hard"] else "Savol yo'q",
                }
                tickets.append(ticket)

        except (IndexError, ValueError):
            return Response({"error": "Tanlash uchun yetarli savollar mavjud emas!"}, status=400)

        # Natijani JSON fayliga saqlash
        tickets_output_path = os.path.join(settings.MEDIA_ROOT, "tickets_output.json")
        with open(tickets_output_path, "w", encoding="utf-8") as f:
            json.dump({"tickets": tickets}, f, ensure_ascii=False, indent=4)

        return Response({
            "message": f"{num_tickets} ta biletlar yaratildi!",
            "files": [f"{settings.MEDIA_URL}tickets_output.json"],
            "tickets": tickets
        })









from datetime import datetime


class ExportTickets(APIView):
    permission_classes = [AllowAny]
    

    def get(self, request):
        pythoncom.CoInitialize()
        bugun = datetime.now()
        formatlangan_sana = bugun.strftime("%Y_%m_%d")

        tickets_output_path = os.path.join(settings.MEDIA_ROOT, "tickets_output.json")
        print(f"Checking tickets JSON at: {tickets_output_path}")

        try:
            with open(tickets_output_path, "r", encoding="utf-8") as f:
                data = json.load(f)
        except FileNotFoundError:
            return Response({"error": "output.json fayli topilmadi!"}, status=404)

        tickets = data.get("tickets", [])
        if not tickets:
            return Response({"error": "Biletlar topilmadi!"}, status=400)

        try:
            samdukf_instance = SamDUkf.objects.latest('id')
            print(f"Using SamDUkf instance: {samdukf_instance}")
        except SamDUkf.DoesNotExist:
            return Response({"error": "SamDUkf ma'lumotlari topilmadi!"}, status=404)

        academic_year = samdukf_instance.uquv_yili
        semester = samdukf_instance.semestr
        subject = samdukf_instance.fan
        stage = samdukf_instance.bosqich
        field_of_study = samdukf_instance.talim_yunalishi
        # talabalar_soni = samdukf_instance.biletlar_soni

        template_path = os.path.join(os.path.dirname(__file__), '../static/bilet_savollar.docx')
        if not os.path.exists(template_path):
            return Response({"error": f"Word shablon fayli topilmadi! Yo'l: {template_path}"}, status=404)

        output_dir = os.path.join(settings.MEDIA_ROOT, "biletlar")
        os.makedirs(output_dir, exist_ok=True)

        word_files = []
        for ticket_num, ticket in enumerate(tickets, start=1):
            doc = Document(template_path)
            if not doc.tables:
                return Response({"error": "Word shablonida jadval mavjud emas!"}, status=500)

            table2 = doc.tables[0]
            cell_1 = table2.rows[0].cells[0]
            cell_1.text = f"{academic_year}-O‘quv yili {stage}-bosqich {semester}-semestr"
            paragraph1 = cell_1.paragraphs[0]
            run = paragraph1.runs[0]
            paragraph1.alignment = 1
            run.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(20)

            cell_2 = table2.rows[1].cells[0]
            cell_2.text = f"{field_of_study} yo‘nalishiga"
            paragraph2 = cell_2.paragraphs[0]
            run = paragraph2.runs[0]
            paragraph2.alignment = 1
            run.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(20)

            cell_3 = table2.rows[2].cells[0]
            cell_3.text = f"{subject} fanidan"
            paragraph3 = cell_3.paragraphs[0]
            run = paragraph3.runs[0]
            paragraph3.alignment = 1
            run.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(20)

            table = doc.tables[-2]
            questions = list(ticket.values())
            for i, question in enumerate(questions):
                if i < len(table.rows):
                    cell = table.rows[i].cells[1]
                    cell.text = ""
                    p = cell.paragraphs[0]
                    run = p.add_run(question)
                    run.font.size = Pt(14)
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            output_file = os.path.join(output_dir, f"bilet_{ticket_num}.docx")
            doc.save(output_file)
            word_files.append(output_file)

        pdf_files = []
        for word_file in word_files:
            pdf_file = word_file.replace(".docx", ".pdf")
            convert(word_file)
            pdf_files.append(pdf_file)

        # Statik fayl nomi bilan merged PDF yaratish
        merged_pdf_path = os.path.join(output_dir, f"{formatlangan_sana}_sanada_{samdukf_instance.biletlar_soni}_ta talabaga_muljallangan_biletlar.pdf")
        merger = PdfMerger()
        for pdf in pdf_files:
            merger.append(pdf)
        merger.write(merged_pdf_path)
        merger.close()
        print(f"Generated merged PDF at: {merged_pdf_path}")

        # Eski faylni ID orqali topib o'chirish
        try:
            # Avvalgi ID li faylni topish
            old_samdukf_doc = SamDUkfDoc.objects.filter(samdukf__id__lt=samdukf_instance.id).last()
            if old_samdukf_doc and old_samdukf_doc.file:
                if os.path.isfile(old_samdukf_doc.file.path):
                    
                    os.remove(old_samdukf_doc.file.path)
                old_samdukf_doc.delete()  # Bazadan ham o'chirish
                
        except Exception as e:
            print(f"Error deleting old file/record: {str(e)}")

        # Yangi faylni saqlash (eski fayl qayta yoziladi)
        try:
            samdukf_doc, created = SamDUkfDoc.objects.get_or_create(samdukf=samdukf_instance)
            with open(merged_pdf_path, 'rb') as pdf_file:
                # Faylni qayta yozish uchun to'g'ridan-to'g'ri save ishlatamiz
                samdukf_doc.file.save(f"{formatlangan_sana}_sanada_{samdukf_instance.biletlar_soni}_ta talabaga_muljallangan_biletlar.pdf", File(pdf_file), save=True)
                
        except Exception as e:
            
            return Response({"error": f"PDF faylni saqlashda xatolik: {str(e)}"}, status=500)

        # Vaqtinchalik fayllarni o'chirish
        for pdf in pdf_files:
            if os.path.exists(pdf):
                try:
                    os.remove(pdf)
                    print(f"Cleaned up temporary PDF: {pdf}")
                except Exception as e:
                    print(f"Failed to clean up temporary PDF: {str(e)}")

        for word in word_files:
            if os.path.exists(word):
                try:
                    os.remove(word)
                    print(f"Cleaned up temporary Word file: {word}")
                except Exception as e:
                    print(f"Failed to clean up temporary Word file: {str(e)}")

        # Agar merged_pdf_path qolsa, uni o'chirish
        if os.path.exists(merged_pdf_path):
            try:
                os.remove(merged_pdf_path)
                print(f"Cleaned up intermediate PDF: {merged_pdf_path}")
            except Exception as e:
                print(f"Failed to clean up intermediate PDF: {str(e)}")

        pythoncom.CoUninitialize()

        return Response({
            "message": "Biletlar PDFga aylantirildi va saqlandi!",
            "merged_pdf": f"{settings.MEDIA_URL}biletlar/merged_tickets_{samdukf_instance.id}.pdf"
        })
    